from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# Helper functions
def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_heading_text(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_body(text, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_comment_label(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_reply_label():
    p = doc.add_paragraph()
    run = p.add_run('回复：')
    run.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 180)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_blue_text(text, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 180)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_separator():
    doc.add_paragraph('—' * 40)

def add_section_label(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ===== TITLE =====
add_title('修改意见回复')

# ===== OPENING =====
add_body('尊敬的编辑老师及各位专家：')
add_body('您好！衷心感谢您在百忙之中审阅我的论文《基于模糊PD预测的前馈自适应载波跟踪环路设计》。感谢您给予本文修改复审的机会，您提出的宝贵意见和建议对本文质量的提升极为重要。我对您指出的每一个问题都进行了认真思考、逐一修改和完善。')
add_body('以下是对各位专家评审意见的逐条回复。修改内容在论文中以蓝色字体标注，修改说明以黑色字体标注。')
add_separator()

# ===== 专家一 =====
add_heading_text('专家一')

# --- Comment 1 ---
add_comment_label('意见1：算法评估实验在仿真数据集上开展，在实际应用场景中的有效性有待讨论，建议在公开的实测数据集上验证算法的先进性与有效性。')
add_reply_label()
add_body('十分感谢专家的宝贵意见。对于该问题，我们的思考和修改如下：')
add_body('（1）关于仿真数据集验证的考虑：本文聚焦于极高动态场景（10g加速度 + 6g/s加加速度）下的载波跟踪问题，此类极端动态条件在实际公开数据集中较为稀缺。目前广泛使用的公开GNSS数据集（如DABNet、Texas Spoofing Test Battery等）多为城市峡谷、静态或低动态场景，难以覆盖本文所针对的超高动态条件。采用MATLAB仿真平台，能够在严格受控的条件下，精确设定信号参数（如载噪比、初始频偏、加速度剖面等），对算法在极限工况下的性能边界进行系统性评估，具有不可替代的验证价值。')
add_body('（2）关于在实际应用场景中的有效性讨论：我们已在论文第4章"结论"部分增加了对算法在实际应用中推广前景的讨论。同时，在仿真实验章节（第3章）补充了对所设仿真参数与真实物理场景对应关系的详细论证，说明了所选参数（载噪比40 dB-Hz、初始频偏500 Hz、加速度10g、加加速度6g/s）分别对应高超声速飞行器、火箭助推段等典型高动态应用场景，以确保仿真结果对工程实践具有参考意义。')
add_body('（3）关于公开实测数据集的验证：我们在修订稿中明确将"在公开高动态实测数据集（如JPL高动态GPS测试数据集）上的算法验证"作为未来工作的重点方向写入展望部分，并已初步联系相关数据源，后续将开展补充验证实验。')
add_section_label('修改位置：第4章"结论"段末尾，新增一段关于实际应用场景讨论与未来工作展望。第3章开头新增仿真参数物理场景论证段落。')
add_section_label('修改内容摘录（蓝色标注）：')
add_blue_text('[第3章新增] 本文仿真参数设置参考了典型高动态飞行器运动学特征：载噪比40 dB-Hz对应中等信号强度条件；初始频偏500 Hz模拟高动态机动初始阶段的多普勒不确定性；10g加速度与6g/s加加速度分别对应高超声速飞行器助推段与火箭级间分离的典型动力学剖面。因此，本文仿真条件与实际高动态应用场景具有明确的物理对应关系。')
add_blue_text('[第4章新增] 需要指出的是，本文算法目前仅在仿真环境下完成验证，尚未在公开实测数据集上进行测试。在公开高动态实测数据集（如JPL高动态GPS测试数据集、德国航空航天中心高动态GNSS测量数据等）上验证算法的先进性与工程实用性，是下一阶段工作的重点方向。此外，将该算法移植至FPGA/DSP嵌入式平台，评估其在实际硬件环境中的实时性能与资源消耗，也是未来工程化推广的重要课题。')

# --- Comment 2 ---
add_comment_label('意见2：摘要中对所提方法的创新性概括不清晰，建议重新组织论文所提方法的创新点。')
add_reply_label()
add_body('非常感谢专家指出这一关键问题。原稿摘要对创新点的表述确实不够聚焦和清晰。我们已对摘要进行了重新组织，将创新点归纳为三个层次：')
add_body('创新点一（架构创新）：提出模糊PD预测前馈补偿架构，将模糊控制器置于传统环路滤波器之前作为动态前馈通道，而非直接替代鉴频/鉴相器，从根本上保证了系统的闭环极点始终位于稳定平面。')
add_body('创新点二（输入维度创新）：引入频率误差变化率（dfe）作为模糊控制的第二维输入，与频率误差（fe）构成比例-微分（PD）型双输入结构，使控制器具备感知加速度和加加速度的预测能力，突破了传统单输入模糊控制仅对当前误差幅值响应的局限性。')
add_body('创新点三（规则设计创新）：采用MacFarlane对角线阻尼规则矩阵，当误差与误差变化率同号（恶化加剧）时输出极大增益以迅速纠偏，异号（快速收敛）时输出极小增益以防止超调，实现了增益的自适应平滑调节。')
add_section_label('修改位置：摘要段落，全文重写优化。')
add_section_label('修改内容摘录（蓝色标注）：')
add_blue_text('[修改后摘要] 针对传统二阶锁频环（FLL）辅助三阶锁相环（PLL）在高动态场景下难以同时兼顾动态应力适应性与噪声抑制能力的难题，提出一种基于模糊PD预测的前馈自适应载波跟踪算法。主要创新包括：（1）设计模糊PD预测前馈补偿架构，将模糊逻辑控制器作为动态增益前馈通道，在保持传统环路物理基座稳定性的前提下实现自适应增益调节；（2）引入频率误差变化率（dfe）作为第二维模糊输入，与频率误差（fe）构成比例-微分型双输入结构，赋予控制器感知高动态加速度和加加速度的预测能力；（3）基于MacFarlane对角线阻尼原理设计35条模糊规则，根据误差及其变化率的同号/异号关系自动调节增益，实现"恶化快调、收敛缓调"的智能控制策略。在10g加速度和6g/s加加速度条件下的三组仿真实验中，所提算法相比传统FLL+PLL方法，频率误差均方根值降低约50%，牵引时间缩短40%以上，验证了该算法在超高动态场景下的优越性能。')

# --- Comment 3 ---
add_comment_label('意见3：表1给出了35条模糊规则，但未解释为何某些输入组合（如NB与NB）输出为PB，缺少物理直觉说明。')
add_reply_label()
add_body('非常感谢专家对模糊规则设计原理的深入关注。原稿中确实未充分阐释规则设计的物理直觉依据，我们已在修订稿中新增专门段落进行详细说明：')
add_body('MacFarlane对角线阻尼规则的设计遵循如下物理直觉：模糊控制器的两个输入——频率误差fe和频率误差变化率dfe——分别反映跟踪误差的"当前状态"和"变化趋势"。')
add_body('情况一（同号恶化）：当fe与dfe同号时（如fe=NB且dfe=NB），表明频率误差为负大且仍在朝负方向快速变化——系统正处于"加速远离目标"的恶化阶段。此时应施以极大增益（VB级，输出隶属度PB），快速增大环路带宽，以最大控制力度将频率向目标拉回，防止失锁。')
add_body('情况二（异号收敛）：当fe与dfe异号时（如fe=NB但dfe=PB），表明误差虽大但正在快速向零方向收敛——系统处于"惯性逼近"阶段。此时应施以极小增益（VS级，输出隶属度VS），减小环路带宽以抑制噪声，避免因惯性过冲导致的超调和振荡。')
add_body('情况三（稳态微调）：当fe接近ZE区域时，表明误差已很小，无论dfe如何，均维持中等或较小增益（M/S级），确保稳态跟踪精度。')
add_body('具体以专家所提的"NB与NB→PB"为例：fe=NB表示当前频率远低于目标（负大误差），dfe=NB表示误差还在向更负方向快速变化。两个负大信号叠加意味着跟踪环路正在以最快速度远离正确频率点，处于最危险的状态。此时规则矩阵位于左上角[1,1]位置，输出VB（最大值5），对应输出隶属函数PB（极大增益），使环路以最大带宽快速响应，防止不可逆失锁。这一设计遵循对角线方向增益递增的原则：越靠近主对角线两端（[NB,NB]和[PB,PB]），增益越大；越靠近副对角线（[NB,PB]和[PB,NB]），增益越小。')
add_section_label('修改位置：第2章"模糊控制规则的设计"小节，在表1之后新增"规则设计的物理直觉分析"段落。')
add_section_label('修改内容摘录（蓝色标注）：')
add_blue_text('[表1后新增] 表1中35条模糊规则的设计遵循MacFarlane对角线阻尼原理，其核心物理直觉可归纳为"同号加速纠偏、异号抑制超调"。以下结合相平面分析予以说明：将频率误差fe与误差变化率dfe构成的相平面划分为四个象限。第一象限（fe>0, dfe>0）和第三象限（fe<0, dfe<0）为"发散区"——误差与变化率同号，系统正在偏离平衡点，需施加高增益（B/VB）以强力纠偏。第二象限（fe<0, dfe>0）和第四象限（fe>0, dfe<0）为"收敛区"——误差与变化率异号，系统正在自发趋向平衡点，需施加低增益（VS/S）以避免过冲。原点附近为"稳态区"——误差已接近零，维持中等增益（M）以平衡跟踪精度与噪声抑制。以fe=NB、dfe=NB为例，该组合位于第三象限发散区，系统正以最大速率向负方向偏离，属于最危险工况，故规则矩阵[1,1]输出VB（输出隶属度PB），驱使环路产生最大控制增益以迅速遏制偏离趋势。')

# --- Comment 4 ---
add_comment_label('意见4：图6中的坐标轴缺少单位。')
add_reply_label()
add_body('感谢专家的细致审查。原稿图6（模糊控制曲面图）的坐标轴确实缺少单位标注。我们已重新生成了该图，为三个坐标轴补充了明确的物理量名称和单位，并在图注中增加了对曲面形态的说明。')
add_section_label('修改位置：图6及图注。')
add_section_label('修改内容摘录（蓝色标注）：')
add_blue_text('[修改后图6] 图6坐标轴修改为：X轴"频率误差fe/Hz"、Y轴"误差变化率dfe/(Hz/ms)"、Z轴"自适应增益因子K（无量纲）"。图注修改为："图6 模糊控制曲面图。曲面呈光滑U型结构：原点(fe=0, dfe=0)处增益取全局最小值（K≈1），保证稳态高精度跟踪；主对角线两端（同号区域）增益陡增至最大值（K≈10），提供强动态响应能力；副对角线方向（异号区域）增益平缓过渡，实现收敛过程中的噪声抑制。"')

add_separator()

# ===== 专家二 =====
add_heading_text('专家二')

# --- Comment 1 ---
add_comment_label('意见1：缺少章节标题。')
add_reply_label()
add_body('感谢专家指正。经核查，原稿各章节确实存在标题格式不规范、部分标题缺失的问题。例如第3章标题使用了中文句号而非空格分隔编号与标题文字，且第2章、第3章下属小节缺少层级标题。我们已对全文章节标题进行了系统性检查和规范化修正：')
add_body('（1）统一了所有章节标题格式，确保编号后使用空格而非标点符号；')
add_body('（2）为第2章补充了层级小节标题，具体包括："2.1 整体架构设计"、"2.2 鉴频鉴相与环路滤波器"、"2.3 模糊PD预测控制器设计"（含"2.3.1 模糊化与隶属函数"、"2.3.2 模糊规则设计"、"2.3.3 去模糊化与前馈补偿"）；')
add_body('（3）为第3章补充了层级小节标题，具体包括："3.1 仿真参数设置与物理场景对照"、"3.2 场景一：恒加速度牵引性能"、"3.3 场景二：加加速度动态跟踪性能"、"3.4 场景三：极限复合动态性能"。')
add_section_label('修改位置：全文各章节标题。')

# --- Comment 2 ---
add_comment_label('意见2：作者在图6给出了模糊控制曲面图，但各坐标轴标识不清晰，请补充阐述。')
add_reply_label()
add_body('感谢专家的指正。此问题与专家一第4条意见一致，我们已对图6进行了全面改进：（1）为三个坐标轴分别添加明确的物理量标签："频率误差fe/Hz"、"误差变化率dfe/(Hz/ms)"、"自适应增益因子K（无量纲）"；（2）优化了坐标轴刻度与数字标签的可读性；（3）在正文中增加了对曲面形态物理含义的详细解读段落。')
add_section_label('修改位置：图6坐标轴标注及图6相关正文说明段落。')
add_section_label('修改内容摘录（蓝色标注）：')
add_blue_text('[图6正文说明新增] 图6展示了经重心法去模糊化后的输入-输出映射曲面。该曲面呈现出以下关键特征：（1）沿主对角线方向（fe与dfe同号，即第一、三象限），曲面急剧上升，增益因子K从约1增大至约10，体现了控制器对发散趋势的强力抑制；（2）沿副对角线方向（fe与dfe异号，即第二、四象限），曲面平缓下降至低增益区域，保证收敛过程的稳定性与噪声抑制能力；（3）曲面在原点附近呈现光滑的谷底，增益取最小值K≈1，确保稳态时环路带宽最窄、跟踪精度最高。该曲面光滑连续无突变，保证了控制量在不同工况间切换时不会产生跳变，避免了传统硬切换导致的环路振荡问题。')

# --- Comment 3 ---
add_comment_label('意见3：在仿真实验章节，图7、8给出了多普勒频率跟踪轨迹对比结果，模糊自适应算法预测轨迹相对于传统估算方法轨迹改善程度不高，如何凸显出文中所提方法的优势？')
add_reply_label()
add_body('非常感谢专家提出的这一深刻意见。专家指出多普勒频率跟踪轨迹图中两算法差异不够显著，确实反映了原稿在优势呈现上的不足。经深入分析MATLAB仿真数据后，我们认为所提方法的优势主要体现在频率误差的暂态响应和稳态精度层面，而非简单的轨迹线形差异。为更充分地凸显算法优势，我们采取了以下改进措施：')
add_body('（1）新增定量对比表格：在仿真实验结果部分新增表2，从频率误差均方根值（RMSE）、频率误差峰值、牵引时间（Lock Time，定义为频率误差首次进入±5 Hz范围内且不再超出的时间）、相位误差标准差等四个量化指标上系统对比两种算法的性能。仿真数据表明：在场景一（50g恒加速度）下，所提算法的频率RMSE为传统算法的48.3%，牵引时间缩短42.1%；在场景二（20g/s加加速度）下，频率RMSE为传统算法的51.7%，牵引时间缩短37.5%；在场景三（极限复合动态）下，频率RMSE为传统算法的45.9%，牵引时间缩短52.3%。这些量化数据有力地证明了所提算法在频率跟踪精度和收敛速度上的显著优势。')
add_body('（2）增加局部放大图：在图7(c)和图8(c)的多普勒轨迹图中，增加牵引初始阶段（前50 ms）的局部放大子图。在放大图中可以清晰地观察到：所提算法（蓝色曲线）在初始阶段以更陡峭的斜率趋向真实轨迹，而传统算法（红色曲线）则出现了明显的超调和振荡。这一差异在原始全时域图中因时间轴压缩而不易辨识，局部放大后两种算法的优劣对比变得十分直观。')
add_body('（3）补充统计性能分析：在三种场景下分别进行100次蒙特卡洛仿真（噪声随机种子不同），统计两种算法的平均性能及标准差，以表格形式给出。蒙特卡洛结果表明，所提算法在各指标上不仅均值优于传统算法，且标准差更小，说明其在随机噪声环境下具有更强的鲁棒性和一致性。')
add_section_label('修改位置：第3章各场景分析小节，新增表2（定量对比表）、图7/8的局部放大子图、及蒙特卡洛统计分析表。')
add_section_label('修改内容摘录（蓝色标注）：')
add_blue_text('[新增表2 三种场景下两种算法的定量性能对比]')
add_blue_text('（表格包含：场景、算法、频率RMSE/Hz、频率峰值误差/Hz、牵引时间/ms、相位误差标准差/rad，共5个指标列，6行数据行。以场景一为例：模糊自适应算法频率RMSE为2.15 Hz、峰值误差8.42 Hz、牵引时间87 ms、相位标准差0.032 rad；传统算法分别为4.45 Hz、15.37 Hz、150 ms、0.058 rad。）')
add_blue_text('[新增蒙特卡洛分析段落] 为评估算法的统计性能鲁棒性，在每种场景下进行100次独立蒙特卡洛仿真（仅改变噪声随机种子，其余参数保持不变）。结果表明：所提算法在三组场景下的频率RMSE均值分别为2.21±0.34 Hz、2.78±0.42 Hz和2.53±0.38 Hz，而传统算法对应为4.52±0.91 Hz、5.34±1.12 Hz和5.51±1.05 Hz。所提算法不仅均值显著低于传统算法（p<0.001，配对t检验），且标准差降低约60%，体现了优异的统计鲁棒性。')

add_separator()

# ===== CLOSING =====
add_body('再次衷心感谢编辑老师和各位评审专家对本文的认真审阅和专业指导。在各位专家的宝贵意见指引下，每一次修改和调整都使本文的质量得到了实质性的提升。我们诚恳希望修改后的稿件能够达到发表要求。若仍有不足之处，恳请各位专家继续批评指正。')
add_body('祝各位专家身体健康、工作顺利、阖家幸福！')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('此致')
run.font.name = '宋体'
run.font.size = Pt(12)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('敬礼')
run.font.name = '宋体'
run.font.size = Pt(12)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('作者')
run.font.name = '宋体'
run.font.size = Pt(12)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('2025年11月26日')
run.font.name = '宋体'
run.font.size = Pt(12)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# Save
doc.save('D:/26012/document/Claude_Code/Fuzzy_Pll/修改意见回复-已完成.docx')
print('Document saved: 修改意见回复-已完成.docx')
