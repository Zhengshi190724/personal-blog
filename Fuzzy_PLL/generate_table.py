from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def set_cell_font(cell, text, bold=False, size=10, color=None):
    """Set cell text with formatting"""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(size)
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if color:
        run.font.color.rgb = color

def set_cell_shading(cell, color_hex):
    """Set cell background color"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)

# ============ TITLE ============
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('修改意见回复 —— 专家二意见3补充：定量性能对比表')
run.bold = True
run.font.size = Pt(14)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()

# ============ 定量分析说明 ============
p = doc.add_paragraph()
run = p.add_run('补充说明：')
run.bold = True
run.font.name = '宋体'
run.font.size = Pt(12)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

explanation = (
    '针对专家二意见3，我们对模糊自适应算法与传统FLL+PLL算法在三种场景下的性能进行了定量对比分析。'
    '评价指标包括频率误差均方根值（RMSE）、频率误差峰值、牵引时间（50ms滑动窗口频率误差RMS首次低于15Hz的时刻）'
    '和相位误差标准差。结果如下表所示，模糊自适应算法在所有指标上均显著优于传统方法。'
)
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0.74)
run = p.add_run(explanation)
run.font.name = '宋体'
run.font.size = Pt(12)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()

# ============ TABLE ============
# 10 rows: 1 header + 3 scenarios × 3 rows (scenario name + 2 algo rows)
table = doc.add_table(rows=10, cols=6, style='Table Grid')

# Header row
headers = ['场景', '算法', '频率RMSE\n(Hz)', '峰值误差\n(Hz)', '牵引时间\n(ms)', '相位标准差\n(rad)']
for j, h in enumerate(headers):
    set_cell_font(table.rows[0].cells[j], h, bold=True, size=9)
    set_cell_shading(table.rows[0].cells[j], '4472C4')
    # Make header text white
    for p in table.rows[0].cells[j].paragraphs:
        for run in p.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

# Data
data = [
    # [row_label, is_scenario_name?, algo_name, rmse, peak, lock, phase_std]
    ['场景一：50g恒加速度\n(初始频偏100Hz)', True, '', '', '', '', ''],
    ['', False, '模糊自适应', '20.23', '104.61', '282.0', '0.9417'],
    ['', False, '传统FLL+PLL', '53.28', '185.80', '445.0', '1.1327'],
    ['场景二：20g/s加加速度\n(初始频偏100Hz)', True, '', '', '', '', ''],
    ['', False, '模糊自适应', '7.79', '90.72', '26.0', '0.6453'],
    ['', False, '传统FLL+PLL', '12.66', '103.28', '86.0', '0.8552'],
    ['场景三：50g加速度\n+20g/s加加速度\n(初始频偏500Hz)', True, '', '', '', '', ''],
    ['', False, '模糊自适应', '40.24', '354.03', '419.0', '0.9964'],
    ['', False, '传统FLL+PLL', '91.98', '456.02', '626.0', '1.1100'],
]

row_idx = 1
for d in data:
    cells = table.rows[row_idx].cells
    if d[1]:  # scenario name row
        # Merge cells 0-5 for scenario name
        merged = cells[0]
        for j in range(1, 6):
            merged = merged.merge(cells[j])
        set_cell_font(table.rows[row_idx].cells[0], d[0], bold=True, size=9)
        set_cell_shading(table.rows[row_idx].cells[0], 'D9E2F3')
    else:
        set_cell_font(cells[0], d[0], size=9)
        set_cell_font(cells[1], d[2], size=9)
        set_cell_font(cells[2], d[3], size=9)
        set_cell_font(cells[3], d[4], size=9)
        set_cell_font(cells[4], d[5], size=9)
        set_cell_font(cells[5], d[6], size=9)
        # Highlight fuzzy row
        if '模糊' in d[2]:
            for j in range(6):
                set_cell_shading(cells[j], 'E2EFDA')
    row_idx += 1

doc.add_paragraph()

# ============ 改善幅度汇总 ============
p = doc.add_paragraph()
run = p.add_run('改善幅度汇总：')
run.bold = True
run.font.name = '宋体'
run.font.size = Pt(12)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

summary_text = (
    '场景一：频率RMSE降低62.0%，峰值误差降低43.7%，牵引时间缩短36.6%，相位标准差降低16.9%\n'
    '场景二：频率RMSE降低38.5%，峰值误差降低12.2%，牵引时间缩短69.8%，相位标准差降低24.6%\n'
    '场景三：频率RMSE降低56.3%，峰值误差降低22.4%，牵引时间缩短33.1%，相位标准差降低10.2%'
)
p = doc.add_paragraph()
run = p.add_run(summary_text)
run.font.name = '宋体'
run.font.size = Pt(11)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()

# ============ 结论 ============
p = doc.add_paragraph()
run = p.add_run('结论：')
run.bold = True
run.font.name = '宋体'
run.font.size = Pt(12)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

conclusion = (
    '定量分析结果表明，模糊自适应算法在全部三种高动态场景下，所有四项性能指标均显著优于传统FLL+PLL方法。'
    '其中，频率RMSE改善幅度达38.5%~62.0%，牵引时间缩短33.1%~69.8%，充分验证了所提模糊PD预测前馈补偿架构'
    '在超高动态载波跟踪中的先进性。这些量化数据有力地回应了专家关于"如何凸显算法优势"的关切。'
)
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0.74)
run = p.add_run(conclusion)
run.font.name = '宋体'
run.font.size = Pt(12)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# Save
doc.save('D:/26012/document/Claude_Code/Fuzzy_Pll/定量性能对比表.docx')
print('Saved: 定量性能对比表.docx')
